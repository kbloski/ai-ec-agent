from infrastructure.logging.logger import Logger
from docx import Document

class DocxParser:
    def __init__(self, logger : Logger ):
        self.logger = logger

    def extract(self, path: str) -> str:
        self.logger.info(f"Opening file: {path}")
        doc = Document(path)

        text = "\n".join(p.text for p in doc.paragraphs)

        self.logger.info(f"Extracted {len(doc.paragraphs)} paragraphs")
        return text